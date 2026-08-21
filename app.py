import sys
import os
import re
import csv
import json
import mimetypes
import urllib.request
import urllib.parse
from html.parser import HTMLParser

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.cell.text import InlineFont
from openpyxl.cell.rich_text import TextBlock, CellRichText

import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QTextEdit, QPushButton, QFileDialog, QProgressBar,
    QMessageBox, QComboBox
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal

def get_base_dir():
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

def get_google_credentials():
    try:
        from google.oauth2.credentials import Credentials
        from google.auth.transport.requests import Request

        base_dir = get_base_dir()
        
        token_path = os.path.join(base_dir, 'token.json')
        if not os.path.exists(token_path):
            token_path = os.path.join(os.getcwd(), 'token.json')

        secret_path = os.path.join(base_dir, 'client_secret.json')
        if not os.path.exists(secret_path):
            secret_path = os.path.join(os.getcwd(), 'client_secret.json')
        
        if os.path.exists(token_path) and os.path.exists(secret_path):
            with open(token_path, 'r', encoding='utf-8') as f:
                token_data = json.load(f)
            with open(secret_path, 'r', encoding='utf-8') as f:
                secret_data = json.load(f)

            installed = secret_data.get('installed') or secret_data.get('web', {})
            client_id = installed.get('client_id')
            client_secret = installed.get('client_secret')

            creds = Credentials(
                token=token_data.get('access_token'),
                refresh_token=token_data.get('refresh_token'),
                token_uri='https://oauth2.googleapis.com/token',
                client_id=client_id,
                client_secret=client_secret
            )

            # Auto-refresh token if expired
            if creds.expired and creds.refresh_token:
                creds.refresh(Request())
                token_data['access_token'] = creds.token
                with open(token_path, 'w', encoding='utf-8') as f:
                    json.dump(token_data, f, indent=2)

            return creds
    except Exception:
        pass
    return None

class HTMLTextExtractor(HTMLParser):
    def __init__(self):
        super().__init__()
        self.result = []
        self.skip = False

    def handle_starttag(self, tag, attrs):
        if tag in ['script', 'style']:
            self.skip = True

    def handle_endtag(self, tag):
        if tag in ['script', 'style']:
            self.skip = False
        if tag in ['p', 'br', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'tr']:
            self.result.append('\n')

    def handle_data(self, data):
        if not self.skip:
            self.result.append(data)

    def get_text(self):
        text = ''.join(self.result)
        lines = [line.strip() for line in text.splitlines()]
        return '\n'.join([line for line in lines if line])

def sanitize_filename(name):
    clean = re.sub(r'[\\/*?:"<>|]', '_', name).strip()
    return clean if clean else "download"

def download_url_content(url, is_binary=False):
    req = urllib.request.Request(
        url, 
        headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
    )
    with urllib.request.urlopen(req) as response:
        info = response.info()
        content_disposition = info.get('Content-Disposition', '')
        content_type = info.get_content_type()
        
        extracted_name = None
        if content_disposition:
            cd_match = re.search(r'filename\*?=(?:UTF-8\'\')?["\']?([^"\';]+)["\']?', content_disposition)
            if cd_match:
                extracted_name = urllib.parse.unquote(cd_match.group(1))

        if is_binary:
            return response.read(), extracted_name, content_type
        return response.read().decode('utf-8', errors='ignore'), extracted_name, content_type

def extract_id_from_url(url, pattern):
    match = re.search(pattern, url)
    return match.group(1) if match else None

def rgb_to_hex(color_dict):
    if not color_dict:
        return None
    r = int(round(color_dict.get('red', 0) * 255))
    g = int(round(color_dict.get('green', 0) * 255))
    b = int(round(color_dict.get('blue', 0) * 255))
    return f"{r:02X}{g:02X}{b:02X}"

def parse_border_side(border_dict):
    if not border_dict or border_dict.get('style') == 'NONE':
        return None
    g_style = border_dict.get('style', '')
    style_map = {
        'DOTTED': 'dotted',
        'DASHED': 'dashed',
        'SOLID': 'thin',
        'SOLID_MEDIUM': 'medium',
        'SOLID_THICK': 'thick',
        'DOUBLE': 'double'
    }
    ox_style = style_map.get(g_style, 'thin')
    color = border_dict.get('color') or border_dict.get('colorStyle', {}).get('rgbColor')
    hex_c = rgb_to_hex(color) or '000000'
    return Side(style=ox_style, color=hex_c)

class BatchDownloadWorker(QThread):
    progress = pyqtSignal(str)
    item_finished = pyqtSignal(str, bool, str)
    all_finished = pyqtSignal(int, int)

    def __init__(self, urls, save_dir, mode="native"):
        super().__init__()
        self.urls = urls
        self.save_dir = save_dir
        self.mode = mode  # "native", "csv", "txt"

    def run(self):
        success_count = 0
        fail_count = 0

        total = len(self.urls)
        for idx, raw_url in enumerate(self.urls, 1):
            url = raw_url.strip()
            if not url:
                continue

            self.progress.emit(f"[{idx}/{total}] Processing: {url}")
            ok, msg = self.process_single_url(url)
            if ok:
                success_count += 1
                self.progress.emit(f"✓ [{idx}/{total}] Success: {msg}\n")
                self.item_finished.emit(url, True, msg)
            else:
                fail_count += 1
                self.progress.emit(f"✗ [{idx}/{total}] Failed: {msg}\n")
                self.item_finished.emit(url, False, msg)

        self.all_finished.emit(success_count, fail_count)

    def process_single_url(self, url):
        try:
            url_lower = url.lower()
            if "trello.com/b/" in url_lower:
                return self.process_trello(url)
            elif "docs.google.com/document" in url_lower:
                return self.process_google_doc(url)
            elif "docs.google.com/spreadsheets" in url_lower:
                return self.process_google_sheet(url)
            elif "drive.google.com/file" in url_lower or "drive.google.com/open" in url_lower:
                return self.process_google_drive_file(url)
            else:
                return False, "Unrecognized URL type. Supported: Google Docs, Google Sheets, Drive files, Trello boards."
        except Exception as e:
            return False, f"Unexpected error: {str(e)}"

    def process_google_doc(self, url):
        doc_id = extract_id_from_url(url, r'/document/d/([a-zA-Z0-9_-]+)')
        if not doc_id:
            return False, "Could not parse Google Doc ID from URL."

        # Plain text mode
        if self.mode in ["txt", "csv"]:
            export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
            try:
                content, _, _ = download_url_content(export_url)
                if not content.startswith("<!DOCTYPE html>"):
                    filename = sanitize_filename(f"GoogleDoc_{doc_id}") + ".txt"
                    filepath = os.path.join(self.save_dir, filename)
                    with open(filepath, "w", encoding="utf-8") as f:
                        f.write(content)
                    return True, f"Extracted Google Doc (.txt) to {filepath}"
            except Exception:
                pass

        # 1:1 Native Mode (.docx)
        if self.mode == "native":
            export_docx_url = f"https://docs.google.com/document/d/{doc_id}/export?format=docx"
            try:
                content, _, _ = download_url_content(export_docx_url, is_binary=True)
                if not content.startswith(b"<!DOCTYPE html>") and len(content) > 500:
                    filename = sanitize_filename(f"GoogleDoc_{doc_id}") + ".docx"
                    filepath = os.path.join(self.save_dir, filename)
                    with open(filepath, "wb") as f:
                        f.write(content)
                    return True, f"Exported 1:1 Google Doc (.docx) to {filepath}"
            except Exception:
                pass

        # Authorized Docs API extraction with High Fidelity
        creds = get_google_credentials()
        if creds:
            try:
                from googleapiclient.discovery import build
                service = build('docs', 'v1', credentials=creds)
                doc_obj = service.documents().get(documentId=doc_id).execute()
                title = doc_obj.get('title', f"GoogleDoc_{doc_id}")

                if self.mode == "native":
                    doc = docx.Document()
                    body_content = doc_obj.get('body', {}).get('content', [])

                    for struct_elem in body_content:
                        paragraph = struct_elem.get('paragraph')
                        if paragraph:
                            named_style = paragraph.get('paragraphStyle', {}).get('namedStyleType', 'NORMAL_TEXT')
                            
                            # Map Headings and Titles
                            docx_style = 'Normal'
                            if named_style == 'TITLE':
                                docx_style = 'Title'
                            elif named_style == 'SUBTITLE':
                                docx_style = 'Subtitle'
                            elif named_style.startswith('HEADING_'):
                                level = named_style.replace('HEADING_', '')
                                docx_style = f'Heading {level}'
                            elif paragraph.get('bullet'):
                                docx_style = 'List Bullet'

                            try:
                                p = doc.add_paragraph(style=docx_style)
                            except Exception:
                                p = doc.add_paragraph()

                            # Alignment
                            alignment = paragraph.get('paragraphStyle', {}).get('alignment')
                            if alignment == 'CENTER':
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif alignment == 'END':
                                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            elif alignment == 'JUSTIFIED':
                                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                            for elem in paragraph.get('elements', []):
                                text_run = elem.get('textRun')
                                if text_run:
                                    run_text = text_run.get('content', '')
                                    run_style = text_run.get('textStyle', {})
                                    
                                    run = p.add_run(run_text)
                                    if run_style.get('bold'):
                                        run.bold = True
                                    if run_style.get('italic'):
                                        run.italic = True
                                    if run_style.get('underline'):
                                        run.underline = True
                                    if run_style.get('strikethrough'):
                                        run.font.strike = True
                                    
                                    font_family = run_style.get('weightedFontFamily', {}).get('fontFamily')
                                    if font_family:
                                        run.font.name = font_family

                                    font_size = run_style.get('fontSize', {}).get('magnitude')
                                    if font_size:
                                        run.font.size = Pt(font_size)

                                    fg_color = run_style.get('foregroundColor', {}).get('color', {}).get('rgbColor')
                                    if fg_color:
                                        r = int(round(fg_color.get('red', 0) * 255))
                                        g = int(round(fg_color.get('green', 0) * 255))
                                        b = int(round(fg_color.get('blue', 0) * 255))
                                        run.font.color.rgb = RGBColor(r, g, b)

                        table = struct_elem.get('table')
                        if table:
                            rows = table.get('tableRows', [])
                            if rows:
                                num_rows = len(rows)
                                num_cols = len(rows[0].get('tableCells', []))
                                docx_table = doc.add_table(rows=num_rows, cols=num_cols)
                                docx_table.style = 'Table Grid'

                                for r_idx, row in enumerate(rows):
                                    for c_idx, cell in enumerate(row.get('tableCells', [])):
                                        cell_paragraphs = []
                                        for c_content in cell.get('content', []):
                                            c_para = c_content.get('paragraph')
                                            if c_para:
                                                para_text = "".join([
                                                    elem.get('textRun', {}).get('content', '')
                                                    for elem in c_para.get('elements', [])
                                                ]).strip()
                                                if para_text:
                                                    cell_paragraphs.append(para_text)
                                        docx_table.cell(r_idx, c_idx).text = "\n".join(cell_paragraphs)

                    filename = sanitize_filename(title) + ".docx"
                    filepath = os.path.join(self.save_dir, filename)
                    doc.save(filepath)
                    return True, f"Created 1:1 Google Doc (.docx) at {filepath}"

                else:
                    # Text extraction
                    text_chunks = []
                    for item in doc_obj.get('body', {}).get('content', []):
                        paragraph = item.get('paragraph')
                        if paragraph:
                            for elem in paragraph.get('elements', []):
                                text_run = elem.get('textRun')
                                if text_run and 'content' in text_run:
                                    text_chunks.append(text_run['content'])
                    full_text = "".join(text_chunks)
                    filename = sanitize_filename(title) + ".txt"
                    filepath = os.path.join(self.save_dir, filename)
                    with open(filepath, "w", encoding="utf-8") as f:
                        f.write(full_text)
                    return True, f"Extracted Google Doc (.txt) to {filepath}"

            except Exception as e:
                return False, f"Google Docs API error: {str(e)}"

        return False, "Google Doc is private or restricted. Ensure token.json authorization is present."

    def process_google_sheet(self, url):
        sheet_id = extract_id_from_url(url, r'/spreadsheets/d/([a-zA-Z0-9_-]+)')
        if not sheet_id:
            return False, "Could not parse Google Sheet ID from URL."

        gid_match = re.search(r'[#&?]gid=([0-9]+)', url)
        target_gid = gid_match.group(1) if gid_match else None

        # 1:1 Native High-Fidelity (.xlsx)
        if self.mode == "native":
            creds = get_google_credentials()
            if creds:
                try:
                    from googleapiclient.discovery import build
                    service = build('sheets', 'v4', credentials=creds)

                    full_meta = service.spreadsheets().get(
                        spreadsheetId=sheet_id,
                        includeGridData=True
                    ).execute()

                    title = full_meta.get('properties', {}).get('title', f"GoogleSheet_{sheet_id}")
                    sheets = full_meta.get('sheets', [])

                    # Build GID to title mapping for internal hyperlinks
                    gid_to_title = {}
                    for s_data in sheets:
                        gid_to_title[str(s_data.get('properties', {}).get('sheetId'))] = re.sub(r'[\\/*?:\\[\\]]', '_', s_data.get('properties', {}).get('title', 'Sheet'))[:31]

                    wb = openpyxl.Workbook()
                    wb.remove(wb.active)

                    for s_data in sheets:
                        props = s_data.get('properties', {})
                        s_title = props.get('title', 'Sheet')
                        clean_sheet_title = re.sub(r'[\\/*?:\\[\\]]', '_', s_title)[:31]
                        ws = wb.create_sheet(title=clean_sheet_title)

                        # Set Tab Color
                        tab_color = props.get('tabColor') or props.get('tabColorStyle', {}).get('rgbColor')
                        hex_tab_color = rgb_to_hex(tab_color)
                        if hex_tab_color:
                            ws.sheet_properties.tabColor = hex_tab_color

                        # Set Gridlines
                        grid_props = props.get('gridProperties', {})
                        ws.views.sheetView[0].showGridLines = not grid_props.get('hideGridlines', False)

                        # Handle Merges
                        merges = s_data.get('merges', [])
                        for m in merges:
                            start_row = m.get('startRowIndex', 0) + 1
                            end_row = m.get('endRowIndex', 0)
                            start_col = m.get('startColumnIndex', 0) + 1
                            end_col = m.get('endColumnIndex', 0)
                            if end_row >= start_row and end_col >= start_col:
                                ws.merge_cells(start_row=start_row, start_column=start_col, end_row=end_row, end_column=end_col)

                        # Handle Grid Data
                        grid_data = s_data.get('data', [])
                        for data_block in grid_data:
                            row_offset = data_block.get('startRow', 0)
                            col_offset = data_block.get('startColumn', 0)
                            
                            # Column widths
                            for col_idx, col_meta in enumerate(data_block.get('columnMetadata', [])):
                                col_letter = get_column_letter(col_offset + col_idx + 1)
                                pixel_size = col_meta.get('pixelSize')
                                if pixel_size:
                                    ws.column_dimensions[col_letter].width = max(round(pixel_size / 7.5, 2), 3)

                            # Rows
                            rows = data_block.get('rowData', [])
                            row_metas = data_block.get('rowMetadata', [])

                            for r_idx, row_obj in enumerate(rows):
                                current_row = row_offset + r_idx + 1
                                
                                # Row height
                                if r_idx < len(row_metas):
                                    row_pixel = row_metas[r_idx].get('pixelSize')
                                    if row_pixel:
                                        ws.row_dimensions[current_row].height = round(row_pixel * 0.75, 2)

                                values = row_obj.get('values', [])
                                for c_idx, cell_obj in enumerate(values):
                                    current_col = col_offset + c_idx + 1
                                    cell = ws.cell(row=current_row, column=current_col)

                                    # Cell Value & Multi-style Rich Text
                                    formatted_val = cell_obj.get('formattedValue')
                                    user_val = cell_obj.get('userEnteredValue', {})
                                    runs = cell_obj.get('textFormatRuns')

                                    if 'formulaValue' in user_val:
                                        cell.value = user_val['formulaValue']
                                    elif runs and formatted_val:
                                        # High-fidelity mixed text formatting (multi-color / multi-size in same cell)
                                        rich_text = CellRichText()
                                        raw_text = str(formatted_val)
                                        for run_idx, run_meta in enumerate(runs):
                                            start = run_meta.get('startIndex', 0)
                                            end = runs[run_idx + 1].get('startIndex', len(raw_text)) if (run_idx + 1 < len(runs)) else len(raw_text)
                                            chunk = raw_text[start:end]
                                            if chunk:
                                                r_fmt = run_meta.get('format', {})
                                                r_fg = r_fmt.get('foregroundColor') or r_fmt.get('foregroundColorStyle', {}).get('rgbColor')
                                                r_hex = rgb_to_hex(r_fg)
                                                
                                                font_f = r_fmt.get('fontFamily', 'Calibri')
                                                if ',' in font_f:
                                                    font_f = font_f.split(',')[0].replace('"', '').strip()

                                                inline_font = InlineFont(
                                                    rFont=font_f,
                                                    sz=r_fmt.get('fontSize', 10),
                                                    b=r_fmt.get('bold', False),
                                                    i=r_fmt.get('italic', False),
                                                    strike=r_fmt.get('strikethrough', False),
                                                    u='single' if r_fmt.get('underline') else None,
                                                    color=r_hex if r_hex else None
                                                )
                                                rich_text.append(TextBlock(inline_font, chunk))
                                        cell.value = rich_text
                                    elif 'numberValue' in user_val:
                                        cell.value = user_val['numberValue']
                                    elif 'boolValue' in user_val:
                                        cell.value = user_val['boolValue']
                                    elif formatted_val is not None:
                                        cell.value = formatted_val
                                    elif 'stringValue' in user_val:
                                        cell.value = user_val['stringValue']

                                    # Effective Styling
                                    eff_format = cell_obj.get('effectiveFormat', {})
                                    
                                    # Background Color
                                    bg = eff_format.get('backgroundColor') or eff_format.get('backgroundColorStyle', {}).get('rgbColor')
                                    hex_bg = rgb_to_hex(bg)
                                    if hex_bg and hex_bg != "FFFFFF":
                                        cell.fill = PatternFill(start_color=hex_bg, end_color=hex_bg, fill_type="solid")

                                    # Overall Font Format
                                    tf = eff_format.get('textFormat', {})
                                    fg = tf.get('foregroundColor') or tf.get('foregroundColorStyle', {}).get('rgbColor')
                                    hex_fg = rgb_to_hex(fg)
                                    
                                    font_family = tf.get('fontFamily', 'Calibri')
                                    if ',' in font_family:
                                        font_family = font_family.split(',')[0].replace('"', '').strip()

                                    cell.font = Font(
                                        name=font_family,
                                        size=tf.get('fontSize', 10),
                                        bold=tf.get('bold', False),
                                        italic=tf.get('italic', False),
                                        strike=tf.get('strikethrough', False),
                                        underline='single' if tf.get('underline') else None,
                                        color=hex_fg if (hex_fg and hex_fg != "000000") else None
                                    )

                                    # Borders
                                    borders = eff_format.get('borders', {})
                                    if borders:
                                        cell.border = Border(
                                            top=parse_border_side(borders.get('top')),
                                            bottom=parse_border_side(borders.get('bottom')),
                                            left=parse_border_side(borders.get('left')),
                                            right=parse_border_side(borders.get('right'))
                                        )

                                    # Alignment & Text Wrapping
                                    h_align = eff_format.get('horizontalAlignment', 'LEFT').lower()
                                    v_align = eff_format.get('verticalAlignment', 'BOTTOM').lower()
                                    cell.alignment = Alignment(
                                        horizontal=h_align if h_align in ['left', 'center', 'right'] else None,
                                        vertical=v_align if v_align in ['top', 'center', 'bottom'] else None,
                                        wrap_text=eff_format.get('wrapStrategy') == 'WRAP'
                                    )

                                    # Number / Date / Currency Formatting
                                    num_fmt = eff_format.get('numberFormat', {})
                                    pattern = num_fmt.get('pattern')
                                    if pattern:
                                        cell.number_format = pattern

                                    # Hyperlink Preservation (External URLs and Internal Sheet Links)
                                    hyperlink = cell_obj.get('hyperlink')
                                    if hyperlink:
                                        if hyperlink.startswith('#gid='):
                                            target_gid_key = hyperlink.replace('#gid=', '')
                                            if target_gid_key in gid_to_title:
                                                cell.hyperlink = f"#'{gid_to_title[target_gid_key]}'!A1"
                                        else:
                                            cell.hyperlink = hyperlink

                    filename = sanitize_filename(title) + ".xlsx"
                    filepath = os.path.join(self.save_dir, filename)
                    wb.save(filepath)
                    return True, f"Exported 1:1 Google Sheet (.xlsx) with {len(sheets)} tabs to {filepath}"

                except Exception as e:
                    return False, f"1:1 Sheets API error: {str(e)}"

        # CSV / Plain Text mode
        creds = get_google_credentials()
        if creds:
            try:
                from googleapiclient.discovery import build
                service = build('sheets', 'v4', credentials=creds)
                sheet_meta = service.spreadsheets().get(spreadsheetId=sheet_id).execute()
                title = sheet_meta.get('properties', {}).get('title', f"GoogleSheet_{sheet_id}")
                sheets = sheet_meta.get('sheets', [])

                target_tab_title = None
                if target_gid:
                    for s in sheets:
                        if str(s.get('properties', {}).get('sheetId')) == str(target_gid):
                            target_tab_title = s.get('properties', {}).get('title')
                            break

                ext = ".txt" if self.mode == "txt" else ".csv"
                if target_tab_title:
                    result = service.spreadsheets().values().get(spreadsheetId=sheet_id, range=target_tab_title).execute()
                    rows = result.get('values', [])

                    filename = sanitize_filename(f"{title}_{target_tab_title}") + ext
                    filepath = os.path.join(self.save_dir, filename)
                    if self.mode == "txt":
                        with open(filepath, 'w', encoding='utf-8') as f:
                            for r in rows:
                                f.write("\t".join([str(c) for c in r]) + "\n")
                    else:
                        with open(filepath, 'w', newline='', encoding='utf-8') as f:
                            csv.writer(f).writerows(rows)
                    return True, f"Extracted tab ({ext}) to {filepath}"
                else:
                    saved_count = 0
                    for s in sheets:
                        tab_name = s.get('properties', {}).get('title')
                        result = service.spreadsheets().values().get(spreadsheetId=sheet_id, range=tab_name).execute()
                        rows = result.get('values', [])

                        filename = sanitize_filename(f"{title}_{tab_name}") + ext
                        filepath = os.path.join(self.save_dir, filename)
                        if self.mode == "txt":
                            with open(filepath, 'w', encoding='utf-8') as f:
                                for r in rows:
                                    f.write("\t".join([str(c) for c in r]) + "\n")
                        else:
                            with open(filepath, 'w', newline='', encoding='utf-8') as f:
                                csv.writer(f).writerows(rows)
                        saved_count += 1

                    return True, f"Exported {saved_count} tabs ({ext}) to {self.save_dir}"

            except Exception as e:
                return False, f"Sheets API fallback error: {str(e)}"

        return False, "Google Sheet is private or requires authorization. Check token.json."

    def process_google_drive_file(self, url):
        file_id = extract_id_from_url(url, r'/file/d/([a-zA-Z0-9_-]+)') or extract_id_from_url(url, r'id=([a-zA-Z0-9_-]+)')
        if not file_id:
            return False, "Could not parse Google Drive File ID."

        # Attempt direct public download
        download_url_file = f"https://drive.google.com/uc?export=download&id={file_id}"
        try:
            data, filename_from_header, ctype = download_url_content(download_url_file, is_binary=True)
            if not data.startswith(b"<!DOCTYPE html>"):
                if not filename_from_header:
                    guessed_ext = mimetypes.guess_extension(ctype) if ctype else None
                    if not guessed_ext or guessed_ext == '.bin':
                        guessed_ext = ".pdf"
                    filename_from_header = f"Drive_File_{file_id}{guessed_ext}"
                
                filepath = os.path.join(self.save_dir, sanitize_filename(filename_from_header))
                with open(filepath, "wb") as f:
                    f.write(data)
                return True, f"Downloaded Drive file to {filepath}"
        except Exception:
            pass

        # OAuth Drive API extraction
        creds = get_google_credentials()
        if creds:
            try:
                from googleapiclient.discovery import build
                from googleapiclient.http import MediaIoBaseDownload
                service = build('drive', 'v3', credentials=creds)
                
                file_meta = service.files().get(fileId=file_id, fields='name, mimeType').execute()
                file_name = file_meta.get('name')
                mime_type = file_meta.get('mimeType', '')

                if not file_name:
                    guessed_ext = mimetypes.guess_extension(mime_type) if mime_type else None
                    if not guessed_ext or guessed_ext == '.bin':
                        guessed_ext = ".pdf"
                    file_name = f"Drive_File_{file_id}{guessed_ext}"

                request = service.files().get_media(fileId=file_id)
                filepath = os.path.join(self.save_dir, sanitize_filename(file_name))
                
                with open(filepath, 'wb') as f:
                    downloader = MediaIoBaseDownload(f, request)
                    done = False
                    while not done:
                        status, done = downloader.next_chunk()
                
                return True, f"Downloaded Drive file via API to {filepath}"
            except Exception as e:
                return False, f"Drive API error: {str(e)}"

        return False, "Drive file is restricted or requires owner authentication."

    def process_trello(self, url):
        board_id = extract_id_from_url(url, r'/b/([a-zA-Z0-9_-]+)')
        if not board_id:
            return False, "Could not parse Trello Board ID."

        json_url = f"https://trello.com/b/{board_id}.json"
        try:
            raw_json, _, _ = download_url_content(json_url)
            data = json.loads(raw_json)

            board_name = data.get('name', 'Trello_Board')
            lists = {l['id']: l['name'] for l in data.get('lists', [])}

            lines = []
            lines.append("=" * 80)
            lines.append(f"TRELLO BOARD: {board_name}")
            lines.append("=" * 80 + "\n")

            cards_by_list = {}
            for card in data.get('cards', []):
                if not card.get('closed'):
                    list_name = lists.get(card.get('idList'), 'Unassigned')
                    cards_by_list.setdefault(list_name, []).append(card)

            for list_name, cards in cards_by_list.items():
                lines.append("#" * 80)
                lines.append(f" LIST: {list_name.upper()}")
                lines.append("#" * 80 + "\n")
                for idx, card in enumerate(cards, 1):
                    lines.append(f"--- [Card {idx}] {card.get('name')} ---")
                    desc = card.get('desc', '').strip()
                    if desc:
                        lines.append(f"Description:\n{desc}\n")
                    
                    checklists = card.get('checklists', [])
                    for cl in checklists:
                        lines.append(f"Checklist: {cl.get('name')}")
                        for ci in cl.get('checkItems', []):
                            state = "X" if ci.get('state') == 'complete' else " "
                            lines.append(f"  [{state}] {ci.get('name')}")
                        lines.append("")
                    lines.append("")

            filename = sanitize_filename(board_name) + ".txt"
            filepath = os.path.join(self.save_dir, filename)
            with open(filepath, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))

            return True, f"Exported Trello Board (.txt) to {filepath}"

        except Exception as e:
            return False, f"Trello Board export failed: {str(e)}"

class GenericDownloaderApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle("Downloader")
        self.resize(580, 480)

        centralWidget = QWidget()
        self.setCentralWidget(centralWidget)
        layout = QVBoxLayout()

        # URL Batch Text Area
        layout.addWidget(QLabel("URLs (one or more per line):"))
        self.urlInput = QTextEdit()
        self.urlInput.setPlaceholderText("Paste one or more URLs here (Google Docs, Google Sheets, Drive files, Trello boards)...")
        self.urlInput.setFixedHeight(85)
        layout.addWidget(self.urlInput)

        # Format Mode Selector
        formatLayout = QHBoxLayout()
        formatLayout.addWidget(QLabel("Export Format:"))
        self.formatCombo = QComboBox()
        self.formatCombo.addItem("1:1 Native (.xlsx for Sheets / .docx for Docs / .txt for Trello)", "native")
        self.formatCombo.addItem("Plain Text (.csv for Sheets / .txt for Docs & Trello)", "csv")
        self.formatCombo.addItem("Plain Text Tab-Delimited (.txt for all)", "txt")
        formatLayout.addWidget(self.formatCombo)
        layout.addLayout(formatLayout)

        # Output Folder Field
        layout.addWidget(QLabel("Save Location:"))
        destLayout = QHBoxLayout()
        self.destInput = QTextEdit()
        self.destInput.setFixedHeight(28)
        self.destInput.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.destInput.setText(os.getcwd())
        destLayout.addWidget(self.destInput)

        browseBtn = QPushButton("Browse")
        browseBtn.clicked.connect(self.browseFolder)
        destLayout.addWidget(browseBtn)
        layout.addLayout(destLayout)

        # Download Button
        self.downloadBtn = QPushButton("Download")
        self.downloadBtn.clicked.connect(self.startDownload)
        layout.addWidget(self.downloadBtn)

        self.progressBar = QProgressBar()
        self.progressBar.setRange(0, 0)
        self.progressBar.hide()
        layout.addWidget(self.progressBar)

        # Status / Log Box
        layout.addWidget(QLabel("Status:"))
        self.logText = QTextEdit()
        self.logText.setReadOnly(True)
        layout.addWidget(self.logText)

        centralWidget.setLayout(layout)

    def browseFolder(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Folder", self.destInput.toPlainText().strip())
        if folder:
            self.destInput.setText(folder)

    def log(self, text):
        self.logText.append(text)

    def startDownload(self):
        raw_text = self.urlInput.toPlainText().strip()
        save_dir = self.destInput.toPlainText().strip()
        mode = self.formatCombo.currentData()

        if not raw_text:
            QMessageBox.warning(self, "Warning", "Please enter at least one URL.")
            return

        urls = [line.strip() for line in raw_text.splitlines() if line.strip()]
        if not urls:
            QMessageBox.warning(self, "Warning", "Please enter valid URLs.")
            return

        if not os.path.exists(save_dir):
            QMessageBox.warning(self, "Warning", "Destination folder does not exist.")
            return

        self.downloadBtn.setEnabled(False)
        self.progressBar.show()
        self.log(f"Starting batch queue ({len(urls)} items)...")

        self.worker = BatchDownloadWorker(urls, save_dir, mode=mode)
        self.worker.progress.connect(self.log)
        self.worker.all_finished.connect(self.onFinished)
        self.worker.start()

    def onFinished(self, success_count, fail_count):
        self.downloadBtn.setEnabled(True)
        self.progressBar.hide()

        summary = f"Batch finished: {success_count} succeeded, {fail_count} failed."
        self.log(f"=== {summary} ===")

        if fail_count == 0:
            QMessageBox.information(self, "Success", f"All {success_count} downloads completed successfully!")
        else:
            QMessageBox.warning(self, "Completed with Warnings", summary)

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = GenericDownloaderApp()
    window.show()
    sys.exit(app.exec())
